"""Engine tests — pure functions, no Django, no Word."""

from __future__ import annotations

import io
from datetime import datetime
from decimal import Decimal

import pytest
from docx import Document

from apps.policy_brief import engine


def _minimal_ctx(**overrides) -> engine.BriefContext:
    base = dict(
        case=engine.CaseInfo(
            case_id="HF_TEST_001",
            title="ARNI vs ACEI test",
            intervention="Sacubitril/Valsartan",
            comparator="Enalapril",
            indication="HFrEF",
            population="Adults",
            setting="Hospital",
            perspective="Rumah Sakit",
            created_by_name="Dr. Test",
            status="Disetujui",
        ),
        cea=None,
        bia=None,
        etd=None,
        cba=None,
        recommendation=None,
        references=[],
        approvals=[],
        generated_at=datetime(2026, 6, 8, 14, 30),
        version=1,
    )
    base.update(overrides)
    return engine.BriefContext(**base)


def _read_docx(data: bytes) -> Document:
    return Document(io.BytesIO(data))


class TestMinimalDoc:
    def test_returns_nonempty_bytes(self):
        ctx = _minimal_ctx()
        out = engine.build_docx_bytes(ctx)
        assert isinstance(out, bytes)
        assert len(out) > 5_000  # DOCX zip overhead alone is >5 KB

    def test_cover_section_contains_case_id(self):
        ctx = _minimal_ctx()
        doc = _read_docx(engine.build_docx_bytes(ctx))
        text = "\n".join(p.text for p in doc.paragraphs)
        text += "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        assert "HF_TEST_001" in text
        assert "ARNI vs ACEI test" in text
        assert "Sacubitril/Valsartan" in text

    def test_no_rec_falls_back_to_placeholder(self):
        ctx = _minimal_ctx()
        doc = _read_docx(engine.build_docx_bytes(ctx))
        body = "\n".join(p.text for p in doc.paragraphs)
        assert "Rekomendasi belum tersedia" in body


class TestRecommendationSection:
    def test_green_traffic_light_renders(self):
        ctx = _minimal_ctx(
            recommendation=engine.RecommendationBlock(
                traffic_light="green",
                traffic_light_label="HIJAU — Adopsi tanpa syarat",
                composite_score=Decimal("85.50"),
                evidence_score=Decimal("90.00"),
                ce_score=Decimal("100.00"),
                budget_score=Decimal("70.00"),
                cba_score=Decimal("100.00"),
                justification_text="Strong evidence + cost-saving.",
                computed_at=datetime(2026, 6, 8, 12, 0),
                algorithm_version="1.0.0",
            )
        )
        doc = _read_docx(engine.build_docx_bytes(ctx))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        all_text += "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        assert "HIJAU" in all_text.upper()
        assert "85.50" in all_text
        assert "Strong evidence" in all_text


class TestCEASection:
    def test_cea_block_renders_values(self):
        ctx = _minimal_ctx(
            cea=engine.CEABlock(
                icer=Decimal("10000000"),
                dominance_label="Cost-effective (aman)",
                cost_drug=Decimal("10000000"),
                cost_comparator=Decimal("5000000"),
                efficacy_drug=Decimal("2.5"),
                efficacy_comparator=Decimal("2.0"),
                wtop_threshold=Decimal("250000000"),
                ce_score=Decimal("100"),
                sensitivity_low_icer=Decimal("8000000"),
                sensitivity_high_icer=Decimal("12000000"),
            )
        )
        doc = _read_docx(engine.build_docx_bytes(ctx))
        all_text = "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        body_text = "\n".join(p.text for p in doc.paragraphs)
        combined = all_text + "\n" + body_text
        assert "Cost-effective" in combined
        assert "Rp 10.000.000" in combined
        assert "sensitivitas" in combined.lower()


class TestBIASection:
    def test_bia_three_year_table_has_four_rows(self):
        # Header + 3 years + cumulative = 5 rows
        ctx = _minimal_ctx(
            bia=engine.BIABlock(
                rows=[
                    engine.BIAYearRow(
                        year=1, patients=1000, cost_impact=Decimal("1_000_000_000"),
                        pct_of_budget=Decimal("10.0"),
                    ),
                    engine.BIAYearRow(
                        year=2, patients=1000, cost_impact=Decimal("2_000_000_000"),
                        pct_of_budget=Decimal("20.0"),
                    ),
                    engine.BIAYearRow(
                        year=3, patients=1000, cost_impact=Decimal("3_000_000_000"),
                        pct_of_budget=Decimal("30.0"),
                    ),
                ],
                cumulative_impact=Decimal("6_000_000_000"),
                cumulative_pct_of_budget=Decimal("60.0"),
                severity_label="Significant (10-50%)",
            )
        )
        doc = _read_docx(engine.build_docx_bytes(ctx))
        bia_tables = [
            t for t in doc.tables
            if any("Periode" in cell.text for row in t.rows for cell in row.cells)
        ]
        assert len(bia_tables) == 1
        bia_table = bia_tables[0]
        assert len(bia_table.rows) == 5  # header + 3 years + cumulative

        body = "\n".join(
            cell.text for row in bia_table.rows for cell in row.cells
        )
        assert "Significant" in "\n".join(p.text for p in doc.paragraphs)
        assert "Rp 6.000.000.000" in body


class TestEtDSection:
    def test_etd_renders_9_domains(self):
        rows = [
            engine.EtDDomainRow(
                code=f"domain_{i}",
                label=f"Domain {i}",
                mean_judgement=Decimal("75"),
                judgement_label="Mungkin ya",
                dominant_certainty="Sedang",
                member_count=2,
            )
            for i in range(1, 10)
        ]
        ctx = _minimal_ctx(
            etd=engine.EtDBlock(
                rows=rows,
                aggregated_evidence_score=Decimal("82.50"),
                domains_filled=9,
                domains_total=9,
            )
        )
        doc = _read_docx(engine.build_docx_bytes(ctx))
        etd_tables = [
            t for t in doc.tables
            if any("Domain" == cell.text or "Domain" in cell.text for row in t.rows for cell in row.cells[:1])
        ]
        assert any(
            len(t.rows) >= 10 for t in etd_tables  # header + 9 domains
        )


class TestReferences:
    def test_references_numbered(self):
        refs = [
            engine.ReferenceRow(
                number=1, citation="A et al.", authors="A et al.",
                year=2014, journal="NEJM", doi_or_url="10.1056/NEJMoa1409077",
            ),
            engine.ReferenceRow(
                number=2, citation="B et al.", authors="B et al.",
                year=2020, journal="Lancet", doi_or_url="",
            ),
        ]
        ctx = _minimal_ctx(references=refs)
        doc = _read_docx(engine.build_docx_bytes(ctx))
        all_paragraph_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[1]" in all_paragraph_text
        assert "[2]" in all_paragraph_text
        assert "NEJM" in all_paragraph_text


class TestApprovalSection:
    def test_signature_row_appears(self):
        ctx = _minimal_ctx(
            approvals=[
                engine.ApprovalRow(
                    decision="approved",
                    decision_label="Disetujui",
                    approver_name="Dr. Ketua",
                    approver_email="ketua@test.local",
                    signed_at=datetime(2026, 6, 8, 14, 30),
                    reason="",
                    ip_address="127.0.0.1",
                )
            ]
        )
        doc = _read_docx(engine.build_docx_bytes(ctx))
        all_text = "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        assert "Disetujui" in all_text
        assert "ketua@test.local" in all_text
        assert "127.0.0.1" in all_text


class TestFormatting:
    @pytest.mark.parametrize(
        "value,expected",
        [
            (Decimal("1000"), "Rp 1.000"),
            (Decimal("10000000"), "Rp 10.000.000"),
            (None, "—"),
        ],
    )
    def test_idr_formatting(self, value, expected):
        assert engine._idr(value) == expected

    def test_date_id_formatting(self):
        d = datetime(2026, 6, 8)
        assert engine._date_id(d) == "8 Juni 2026"
