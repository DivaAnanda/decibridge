"""Policy-brief DOCX builder. Pure functions, no Django imports.

Trivially testable: feed in a `BriefContext` dataclass, get bytes out.

Sections (in order):
    1. Cover            — case_id, title, drugs, decision date, signing Ketua name
    2. Executive summary — traffic-light box, composite score, justification
    3. CEA               — ICER, dominance, sensitivity table
    4. BIA               — per-year + cumulative table, % budget
    5. EtD               — per-domain table with mean judgement + dominant certainty
    6. CBA               — table, satisfaction status
    7. References        — numbered bibliography
    8. Audit summary     — signature timestamp + IP + approver name
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from datetime import datetime
from decimal import Decimal

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


# ----------------------------------------------------------------------------
# Data shapes — populated by the service layer from Django models.
# ----------------------------------------------------------------------------


@dataclass(frozen=True)
class CaseInfo:
    case_id: str
    title: str
    intervention: str
    comparator: str
    indication: str
    population: str
    setting: str
    perspective: str
    created_by_name: str
    status: str  # localized label, e.g. "Disetujui"


@dataclass(frozen=True)
class CEABlock:
    icer: Decimal | None
    dominance_label: str  # localized
    cost_drug: Decimal
    cost_comparator: Decimal
    efficacy_drug: Decimal
    efficacy_comparator: Decimal
    wtop_threshold: Decimal
    ce_score: Decimal  # 0-100
    sensitivity_low_icer: Decimal | None = None
    sensitivity_high_icer: Decimal | None = None


@dataclass(frozen=True)
class BIAYearRow:
    year: int
    patients: int
    cost_impact: Decimal
    pct_of_budget: Decimal


@dataclass(frozen=True)
class BIABlock:
    rows: list[BIAYearRow]
    cumulative_impact: Decimal
    cumulative_pct_of_budget: Decimal
    severity_label: str  # localized


@dataclass(frozen=True)
class EtDDomainRow:
    code: str
    label: str  # localized
    mean_judgement: Decimal | None  # 0-100
    judgement_label: str  # localized
    dominant_certainty: str  # localized, e.g. "Tinggi"
    member_count: int


@dataclass(frozen=True)
class EtDBlock:
    rows: list[EtDDomainRow]
    aggregated_evidence_score: Decimal
    domains_filled: int
    domains_total: int


@dataclass(frozen=True)
class CBACriterionRow:
    label: str
    operator: str
    value: str
    satisfied: bool
    notes: str


@dataclass(frozen=True)
class CBABlock:
    rows: list[CBACriterionRow]
    cba_score: Decimal
    satisfied_count: int
    total_count: int


@dataclass(frozen=True)
class RecommendationBlock:
    traffic_light: str  # "green" | "yellow" | "red"
    traffic_light_label: str  # localized
    composite_score: Decimal
    evidence_score: Decimal
    ce_score: Decimal
    budget_score: Decimal
    cba_score: Decimal
    justification_text: str
    computed_at: datetime
    algorithm_version: str


@dataclass(frozen=True)
class ReferenceRow:
    number: int
    citation: str
    authors: str
    year: int | None
    journal: str
    doi_or_url: str


@dataclass(frozen=True)
class ApprovalRow:
    decision: str  # "approved" | "rejected" | "revision_requested"
    decision_label: str  # localized
    approver_name: str
    approver_email: str
    signed_at: datetime
    reason: str
    ip_address: str


@dataclass(frozen=True)
class BriefContext:
    case: CaseInfo
    cea: CEABlock | None
    bia: BIABlock | None
    etd: EtDBlock | None
    cba: CBABlock | None
    recommendation: RecommendationBlock | None
    references: list[ReferenceRow] = field(default_factory=list)
    approvals: list[ApprovalRow] = field(default_factory=list)
    generated_at: datetime = field(default_factory=datetime.now)
    version: int = 1


# ----------------------------------------------------------------------------
# Formatting helpers
# ----------------------------------------------------------------------------


def _idr(value: Decimal | int | float | None) -> str:
    if value is None:
        return "—"
    return f"Rp {value:,.0f}".replace(",", ".")


def _pct(value: Decimal | float | None, decimals: int = 2) -> str:
    if value is None:
        return "—"
    return f"{value:.{decimals}f}%"


def _num(value: Decimal | float | int | None, decimals: int = 2) -> str:
    if value is None:
        return "—"
    if isinstance(value, int):
        return str(value)
    return f"{value:.{decimals}f}"


def _date_id(value: datetime) -> str:
    bulan = [
        "",
        "Januari",
        "Februari",
        "Maret",
        "April",
        "Mei",
        "Juni",
        "Juli",
        "Agustus",
        "September",
        "Oktober",
        "November",
        "Desember",
    ]
    return f"{value.day} {bulan[value.month]} {value.year}"


def _datetime_id(value: datetime) -> str:
    return f"{_date_id(value)}, {value.strftime('%H:%M')} WIB"


TRAFFIC_LIGHT_COLORS = {
    "green": RGBColor(0x2F, 0x9E, 0x44),
    "yellow": RGBColor(0xF0, 0x8C, 0x00),
    "red": RGBColor(0xE0, 0x39, 0x31),
}


# ----------------------------------------------------------------------------
# Document construction
# ----------------------------------------------------------------------------


def _add_heading(doc, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1C, 0x4E, 0x80)


def _add_table(doc, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    head_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        head_cells[i].text = h
        for paragraph in head_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
        head_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for i, val in enumerate(row):
            cells[i].text = val
    return table


def _add_traffic_light_box(doc, recommendation: RecommendationBlock) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Cm(16)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(recommendation.traffic_light_label.upper())
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    bg_hex = {
        "green": "2F9E44",
        "yellow": "F08C00",
        "red": "E03131",
    }.get(recommendation.traffic_light, "808080")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), bg_hex)
    cell._tc.get_or_add_tcPr().append(shading)


# ----------------------------------------------------------------------------
# Section builders
# ----------------------------------------------------------------------------


def _build_cover(doc, ctx: BriefContext) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("RINGKASAN KEBIJAKAN FORMULARIUM")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0x1C, 0x4E, 0x80)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Komite Farmasi dan Terapi (KFT)")
    sub_run.italic = True
    sub_run.font.size = Pt(12)

    doc.add_paragraph()

    info_table = doc.add_table(rows=8, cols=2)
    info_table.style = "Light List Accent 1"
    info = [
        ("ID Kasus", ctx.case.case_id),
        ("Judul", ctx.case.title),
        ("Intervensi", ctx.case.intervention),
        ("Komparator", ctx.case.comparator),
        ("Indikasi", ctx.case.indication),
        ("Status saat ini", ctx.case.status),
        ("Versi dokumen", f"v{ctx.version}"),
        ("Tanggal terbit", _date_id(ctx.generated_at)),
    ]
    for i, (k, v) in enumerate(info):
        cells = info_table.rows[i].cells
        cells[0].text = k
        cells[1].text = v
        for paragraph in cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph()


def _build_executive_summary(doc, ctx: BriefContext) -> None:
    _add_heading(doc, "1. Ringkasan Eksekutif", level=1)
    if not ctx.recommendation:
        doc.add_paragraph("Rekomendasi belum tersedia untuk kasus ini.")
        return

    _add_traffic_light_box(doc, ctx.recommendation)
    doc.add_paragraph()

    summary_rows = [
        ["Skor Komposit", f"{ctx.recommendation.composite_score:.2f} / 100"],
        ["Bukti EtD (bobot 40%)", f"{ctx.recommendation.evidence_score:.2f}"],
        ["Cost-effectiveness (bobot 30%)", f"{ctx.recommendation.ce_score:.2f}"],
        ["Dampak Anggaran (bobot 20%)", f"{ctx.recommendation.budget_score:.2f}"],
        ["Kriteria CBA (bobot 10%)", f"{ctx.recommendation.cba_score:.2f}"],
        ["Algoritma", ctx.recommendation.algorithm_version],
        ["Waktu komputasi", _datetime_id(ctx.recommendation.computed_at)],
    ]
    _add_table(doc, ["Komponen", "Nilai"], summary_rows)

    doc.add_paragraph()
    just_h = doc.add_paragraph()
    just_h.add_run("Justifikasi: ").bold = True
    doc.add_paragraph(ctx.recommendation.justification_text or "—")


def _build_cea(doc, ctx: BriefContext) -> None:
    _add_heading(doc, "2. Analisis Efektivitas Biaya (CEA)", level=1)
    if not ctx.cea:
        doc.add_paragraph("Tidak ada hasil CEA untuk kasus ini.")
        return

    rows = [
        ["ICER", _idr(ctx.cea.icer) + " per unit efektivitas"],
        ["Klasifikasi dominansi", ctx.cea.dominance_label],
        ["Biaya intervensi", _idr(ctx.cea.cost_drug)],
        ["Biaya komparator", _idr(ctx.cea.cost_comparator)],
        ["Efektivitas intervensi (QALY)", _num(ctx.cea.efficacy_drug)],
        ["Efektivitas komparator (QALY)", _num(ctx.cea.efficacy_comparator)],
        ["Ambang WTOP", _idr(ctx.cea.wtop_threshold) + " / QALY"],
        ["Skor CE", f"{ctx.cea.ce_score:.2f} / 100"],
    ]
    _add_table(doc, ["Komponen", "Nilai"], rows)

    if ctx.cea.sensitivity_low_icer is not None and ctx.cea.sensitivity_high_icer is not None:
        doc.add_paragraph()
        sens_p = doc.add_paragraph()
        sens_p.add_run("Analisis sensitivitas ±20%: ").bold = True
        sens_p.add_run(
            f"ICER batas bawah {_idr(ctx.cea.sensitivity_low_icer)} – "
            f"batas atas {_idr(ctx.cea.sensitivity_high_icer)}."
        )


def _build_bia(doc, ctx: BriefContext) -> None:
    _add_heading(doc, "3. Analisis Dampak Anggaran (BIA)", level=1)
    if not ctx.bia:
        doc.add_paragraph("Tidak ada hasil BIA untuk kasus ini.")
        return

    rows: list[list[str]] = []
    for r in ctx.bia.rows:
        rows.append(
            [
                f"Tahun {r.year}",
                f"{r.patients:,}".replace(",", "."),
                _idr(r.cost_impact),
                _pct(r.pct_of_budget),
            ]
        )
    rows.append(
        [
            "Kumulatif",
            "—",
            _idr(ctx.bia.cumulative_impact),
            _pct(ctx.bia.cumulative_pct_of_budget),
        ]
    )
    _add_table(doc, ["Periode", "Pasien", "Dampak Biaya", "% Anggaran"], rows)

    doc.add_paragraph()
    sev_p = doc.add_paragraph()
    sev_p.add_run("Klasifikasi severity: ").bold = True
    sev_p.add_run(ctx.bia.severity_label)


def _build_etd(doc, ctx: BriefContext) -> None:
    _add_heading(doc, "4. Evidence-to-Decision (EtD) — 9 Domain GRADE", level=1)
    if not ctx.etd:
        doc.add_paragraph("Belum ada appraisal EtD untuk kasus ini.")
        return

    rows = []
    for r in ctx.etd.rows:
        rows.append(
            [
                r.label,
                _num(r.mean_judgement, decimals=2),
                r.judgement_label,
                r.dominant_certainty,
                str(r.member_count),
            ]
        )
    _add_table(
        doc,
        ["Domain", "Skor rata-rata", "Judgement", "Certainty dominan", "Voter"],
        rows,
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        f"Domain terisi: {ctx.etd.domains_filled} / {ctx.etd.domains_total} · "
        f"Skor kekuatan bukti agregat: {ctx.etd.aggregated_evidence_score:.2f} / 100"
    )


def _build_cba(doc, ctx: BriefContext) -> None:
    _add_heading(doc, "5. Kriteria Akses (CBA)", level=1)
    if not ctx.cba or not ctx.cba.rows:
        doc.add_paragraph("Tidak ada kriteria CBA yang didefinisikan.")
        return

    rows = []
    for r in ctx.cba.rows:
        rows.append(
            [
                r.label,
                r.operator,
                r.value,
                "Ya" if r.satisfied else "Belum",
                r.notes or "—",
            ]
        )
    _add_table(doc, ["Kriteria", "Operator", "Nilai", "Terpenuhi", "Catatan"], rows)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        f"Status keseluruhan: {ctx.cba.satisfied_count} / {ctx.cba.total_count} kriteria terpenuhi · "
        f"Skor CBA: {ctx.cba.cba_score:.2f} / 100"
    )


def _build_references(doc, ctx: BriefContext) -> None:
    _add_heading(doc, "6. Daftar Referensi", level=1)
    if not ctx.references:
        doc.add_paragraph("Tidak ada referensi terdaftar.")
        return

    for ref in ctx.references:
        parts = []
        if ref.authors:
            parts.append(ref.authors)
        if ref.year:
            parts.append(f"({ref.year})")
        if ref.citation:
            parts.append(ref.citation)
        if ref.journal:
            parts.append(f"<i>{ref.journal}</i>")
        if ref.doi_or_url:
            parts.append(ref.doi_or_url)
        text = f"[{ref.number}] " + ". ".join(p for p in parts if p)
        # Strip <i> placeholder — we add the journal as plain text.
        text = text.replace("<i>", "").replace("</i>", "")
        doc.add_paragraph(text, style="List Number")


def _build_audit_summary(doc, ctx: BriefContext) -> None:
    _add_heading(doc, "7. Ringkasan Audit & Tanda Tangan", level=1)
    if not ctx.approvals:
        doc.add_paragraph("Belum ada tanda tangan tercatat untuk kasus ini.")
        return

    rows = []
    for a in ctx.approvals:
        rows.append(
            [
                _datetime_id(a.signed_at),
                a.decision_label,
                f"{a.approver_name}\n<{a.approver_email}>",
                a.reason or "—",
                a.ip_address or "—",
            ]
        )
    _add_table(doc, ["Waktu", "Keputusan", "Penandatangan", "Alasan", "IP"], rows)

    doc.add_paragraph()
    doc.add_paragraph(
        "Dokumen ini diterbitkan oleh sistem DeciBridge dan tercatat dalam audit log "
        "yang bersifat append-only. Setiap perubahan setelah penerbitan akan tercatat "
        "sebagai versi baru dengan hash SHA-256 berbeda."
    )


def _build_footer(doc, ctx: BriefContext) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"— Dokumen DeciBridge v{ctx.version} · "
        f"diterbitkan {_datetime_id(ctx.generated_at)} —"
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)


# ----------------------------------------------------------------------------
# Public API
# ----------------------------------------------------------------------------


def build_docx_bytes(ctx: BriefContext) -> bytes:
    """Render the policy brief to a DOCX byte string. Pure — no IO."""
    doc = Document()

    # Page setup — A4 portrait, 2.5 cm margins.
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    _build_cover(doc, ctx)
    doc.add_page_break()
    _build_executive_summary(doc, ctx)
    _build_cea(doc, ctx)
    _build_bia(doc, ctx)
    _build_etd(doc, ctx)
    _build_cba(doc, ctx)
    _build_references(doc, ctx)
    _build_audit_summary(doc, ctx)
    _build_footer(doc, ctx)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


__all__ = [
    "ApprovalRow",
    "BIABlock",
    "BIAYearRow",
    "BriefContext",
    "CBABlock",
    "CBACriterionRow",
    "CaseInfo",
    "CEABlock",
    "EtDBlock",
    "EtDDomainRow",
    "RecommendationBlock",
    "ReferenceRow",
    "build_docx_bytes",
]
